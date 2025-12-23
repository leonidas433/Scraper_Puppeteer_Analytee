"""
DOCX Generator Module
Generación de reportes profesionales en formato DOCX
"""

import logging
from pathlib import Path
from datetime import datetime
import re
from typing import List, Optional

logger = logging.getLogger("docx_generator")

class DOCXGeneratorError(Exception):
    pass

class DOCXGenerator:
    """Genera reportes DOCX profesionales con análisis de reseñas"""

    # Colores de marca
    PRIMARY_COLOR = (0x00, 0x4E, 0x89)  # Azul corporativo
    TEXT_COLOR = (0x2C, 0x2C, 0x2C)    # Gris oscuro
    ACCENT_COLOR = (0xE5, 0x3A, 0x3A)  # Rojo para alertas

    def __init__(self, client_name: str, output_dir: str = "./reports"):
        """
        Args:
            client_name: Nombre del cliente/negocio
            output_dir: Directorio de salida
        """
        self.client_name = client_name
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.doc = None

    def create(self, kpis: dict, ai_analysis: str, charts: List[str] = None) -> str:
        """
        Crea documento DOCX con análisis

        Args:
            kpis: Diccionario con KPIs calculados
            ai_analysis: Texto del análisis IA
            charts: Lista de rutas a archivos de gráficos PNG

        Returns:
            Ruta al archivo DOCX generado
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor, Inches, Cm

            logger.info(f"📄 Creando DOCX para {self.client_name}...")

            self.doc = Document()
            self._apply_styles()

            # Portada
            self._add_title_page(kpis)
            self.doc.add_page_break()

            # Headers y footers
            self._add_headers_footers()
            self.doc.add_page_break()

            # Tabla de contenido
            self._add_table_of_contents()
            self.doc.add_page_break()

            # KPIs principales
            self._add_kpi_section(kpis)
            self.doc.add_page_break()

            # Gráficos
            if charts:
                self._add_charts_section(charts)
                self.doc.add_page_break()

            # Análisis IA
            self._add_analysis_section(ai_analysis)
            self.doc.add_page_break()

            # Recomendaciones
            self._add_recommendations_section(kpis)

            # Guardar
            output_file = self._save_document()
            logger.info(f"✅ DOCX generado: {output_file}")

            return output_file

        except ImportError:
            raise DOCXGeneratorError("python-docx no instalada: pip install python-docx")
        except Exception as e:
            logger.error(f"❌ Error generando DOCX: {e}")
            raise DOCXGeneratorError(f"Error: {e}")

    def _apply_styles(self):
        """Aplica estilos globales al documento"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.section import WD_ORIENT, WD_SECTION
        from docx.enum.style import WD_STYLE_TYPE

        # Estilos globales
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(*self.TEXT_COLOR)

        # Márgenes de página
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Estilo de heading personalizado
        if 'CustomHeading1' not in [s.name for s in self.doc.styles]:
            heading1_style = self.doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = 'Calibri'
            heading1_style.font.size = Pt(18)
            heading1_style.font.bold = True
            heading1_style.font.color.rgb = RGBColor(*self.PRIMARY_COLOR)
            heading1_style.paragraph_format.space_after = Pt(12)

        # Estilo de tabla profesional
        if 'ProfessionalTable' not in [s.name for s in self.doc.styles]:
            table_style = self.doc.styles.add_style('ProfessionalTable', WD_STYLE_TYPE.TABLE)
            table_style.font.name = 'Calibri'
            table_style.font.size = Pt(10)

    def _add_title_page(self, kpis):
        """Portada del documento"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        # Título
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("INFORME ORM Y CX")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.PRIMARY_COLOR)

        # Cliente
        client_p = self.doc.add_paragraph()
        client_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_run = client_p.add_run(self.client_name)
        client_run.font.size = Pt(18)

        # Fecha
        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_p.add_run(f"Generado: {datetime.now().strftime('%d/%m/%Y')}")
        date_run.font.size = Pt(10)
        date_run.font.italic = True

        # Métricas principales
        self.doc.add_paragraph()
        avg = kpis.get('avg_rating', 0)
        total = kpis.get('total_reviews', 0)
        nps = kpis.get('nps_estimate', 0)

        metrics = self.doc.add_paragraph()
        metrics.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metrics.add_run(f"⭐ {avg} de 5  |  📊 {total} reseñas  |  📈 NPS {nps}").font.size = Pt(14)

    def _add_headers_footers(self):
        """Agrega headers y footers profesionales"""
        from docx.shared import Pt, RGBColor

        # Header
        header = self.doc.sections[0].header
        header_para = header.paragraphs[0]
        header_run = header_para.add_run(f"Informe ORM - {self.client_name}")
        header_run.font.size = Pt(8)
        header_run.font.color.rgb = RGBColor(*self.PRIMARY_COLOR)
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Footer
        footer = self.doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_run = footer_para.add_run("Generado por Sistema de Análisis de Reseñas")
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Número de página
        footer.add_paragraph("Página ").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_table_of_contents(self):
        """Tabla de contenidos"""
        from docx.shared import Pt

        heading = self.doc.add_paragraph("Tabla de Contenidos", style='CustomHeading1')

        contents = [
            "1. Resumen Ejecutivo",
            "2. Indicadores Clave (KPIs)",
            "3. Análisis Visual (Gráficos)",
            "4. Análisis Detallado",
            "5. Recomendaciones y Plan de Acción"
        ]

        for item in contents:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Pt(18)

    def _add_kpi_section(self, kpis):
        """Sección de KPIs con formato Excel-like"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT

        heading = self.doc.add_paragraph("Indicadores Clave (KPIs)", style='CustomHeading1')

        # Crear tabla profesional con headers
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'ProfessionalTable'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Indicador'
        hdr_cells[1].text = 'Valor'

        # Estilo headers
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(*self.PRIMARY_COLOR)

        # Datos KPIs
        kpi_items = [
            ("Total de Reseñas", str(kpis.get('total_reviews', 0))),
            ("Valoración Media", f"{kpis.get('avg_rating', 0)} ⭐"),
            ("NPS Estimado", f"{kpis.get('nps_estimate', 0)}"),
            ("Frecuencia", f"{kpis.get('review_frequency', {}).get('per_month', 0)} reseñas/mes"),
            ("Sentimiento Positivo", f"{kpis.get('sentiment_distribution', {}).get('positivo', 0)}%"),
        ]

        for i, (label, value) in enumerate(kpi_items, 1):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        # Ajustar ancho de columnas
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(2)

        # Espaciado después de tabla
        self.doc.add_paragraph()

    def _add_charts_section(self, chart_paths: List[str]):
        """Sección de gráficos"""
        from docx.shared import Inches

        heading = self.doc.add_paragraph("Análisis Visual", style='CustomHeading1')

        for chart_path in chart_paths:
            if Path(chart_path).exists():
                self.doc.add_picture(chart_path, width=Inches(5))
                self.doc.add_paragraph()

    def _add_analysis_section(self, ai_analysis: str):
        """Sección con análisis IA mejorado"""
        from docx.shared import Pt, RGBColor

        heading = self.doc.add_paragraph("Análisis Detallado", style='CustomHeading1')

        # Procesar análisis línea por línea
        lines = ai_analysis.split('\n')
        current_table = []
        in_table = False

        for line in lines:
            line = line.strip()
            if not line:
                if in_table:
                    self._process_table(current_table)
                    current_table = []
                    in_table = False
                continue

            # Detectar headings
            if line.startswith('##'):
                if in_table:
                    self._process_table(current_table)
                    current_table = []
                    in_table = False
                self.doc.add_paragraph(line.replace('##', '').strip(), style='Heading 2')
            elif line.startswith('###'):
                if in_table:
                    self._process_table(current_table)
                    current_table = []
                    in_table = False
                self.doc.add_paragraph(line.replace('###', '').strip(), style='Heading 3')
            elif line.startswith('|') and '|' in line:
                # Inicio o continuación de tabla
                current_table.append(line)
                in_table = True
            else:
                if in_table:
                    # Fin de tabla
                    self._process_table(current_table)
                    current_table = []
                    in_table = False
                self.doc.add_paragraph(line)

        # Procesar tabla final si existe
        if in_table and current_table:
            self._process_table(current_table)

    def _process_table(self, table_lines: list):
        """Procesa líneas de tabla Markdown y crea tabla DOCX con formato Excel-like real"""
        if not table_lines:
            return

        # Parsear tabla Markdown correctamente
        rows = []
        for line in table_lines:
            if line.startswith('|') and '|' in line:
                # Dividir por | y limpiar espacios
                cells = [cell.strip() for cell in line.split('|')]
                # Remover celdas vacías al inicio y final
                if cells and not cells[0]: cells = cells[1:]
                if cells and not cells[-1]: cells = cells[:-1]
                if cells:
                    rows.append(cells)

        if not rows or len(rows) < 2:  # Necesitamos al menos header + 1 fila
            return

        # Crear tabla DOCX con estilo Excel-like
        table = self.doc.add_table(rows=len(rows)-1, cols=len(rows[0]))  # -1 porque primera fila es header
        table.style = 'Table Grid'  # Estilo nativo de Word que se ve como Excel

        # Aplicar formato Excel-like completo
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Primera fila como header (formato Excel)
        header_row = rows[0]
        for j, cell_text in enumerate(header_row):
            if j < len(table.rows[0].cells):
                table.rows[0].cells[j].text = cell_text
                # Aplicar estilo header Excel-like
                for paragraph in table.rows[0].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(*self.PRIMARY_COLOR)
                        run.font.size = Pt(10)

        # Filas de datos
        for i, row in enumerate(rows[1:], 1):
            for j, cell_text in enumerate(row):
                if j < len(table.rows[i-1].cells):
                    table.rows[i-1].cells[j].text = cell_text

        # Ajustar anchos de columna proporcionalmente
        total_width = 6.0  # Ancho total en pulgadas
        num_cols = len(rows[0])
        col_width = total_width / num_cols

        for column in table.columns:
            column.width = Inches(col_width)

        self.doc.add_paragraph()  # Espaciado después de tabla

    def _add_recommendations_section(self, kpis):
        """Sección de recomendaciones con mejor formato visual"""
        from docx.shared import Pt, RGBColor

        heading = self.doc.add_paragraph("Recomendaciones y Plan de Acción", style='CustomHeading1')

        # Recomendaciones basadas en KPIs
        avg_rating = kpis.get('avg_rating', 0)
        nps = kpis.get('nps_estimate', 0)

        recommendations = []

        if avg_rating < 3:
            rec = self.doc.add_paragraph(style='List Bullet')
            run = rec.add_run("🔴 CRÍTICO: Mejorar calidad y servicio - Rating muy bajo")
            run.font.color.rgb = RGBColor(220, 53, 69)  # Rojo
            run.font.bold = True
        elif avg_rating < 4:
            rec = self.doc.add_paragraph(style='List Bullet')
            run = rec.add_run("🟠 MODERADO: Revisar puntos débiles identificados")
            run.font.color.rgb = RGBColor(255, 193, 7)  # Amarillo
            run.font.bold = True
        else:
            rec = self.doc.add_paragraph(style='List Bullet')
            run = rec.add_run("🟢 SÓLIDO: Mantener estándares y capitalizar fortalezas")
            run.font.color.rgb = RGBColor(40, 167, 69)  # Verde
            run.font.bold = True

        if nps < 0:
            rec = self.doc.add_paragraph(style='List Bullet')
            run = rec.add_run("Implementar plan de recuperación de clientes insatisfechos")
            run.font.color.rgb = RGBColor(220, 53, 69)
            run.font.bold = True

        # Agregar espaciado final
        self.doc.add_paragraph()

    def _save_document(self) -> str:
        """Guarda el documento"""
        safe_name = re.sub(r'[^A-Za-z0-9_-]', '', self.client_name.replace(' ', '_'))
        filename = f"{safe_name}_Informe_ORM_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self.output_dir / filename

        self.doc.save(str(filepath))
        return str(filepath)